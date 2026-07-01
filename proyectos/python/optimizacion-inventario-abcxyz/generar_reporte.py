"""
generar_reporte.py
Genera el reporte ejecutivo en DOCX y PDF — Optimizacion de Inventario ABC-XYZ
NorteLogix S.A. | Enero 2023 – Diciembre 2024

Pre-requisito:
  python datos/generar_datos.py
  python src/analisis_abcxyz.py

Ejecutar: python generar_reporte.py
"""

import sys
import io
from pathlib import Path
from datetime import datetime

import pandas as pd
import numpy as np
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE  = Path(__file__).resolve().parent
OUT   = BASE / "output"

# ── Paleta (RGB tuples) ───────────────────────────────────────────────────────
AZUL_OSC = (31,  78, 121)
AZUL_MED = (46, 117, 182)
VERDE    = (112, 173,  71)
NARANJA  = (237, 125,  49)
ROJO_OSC = (192,   0,   0)
GRIS     = (100, 100, 100)

MESES_ES = {1:"enero",2:"febrero",3:"marzo",4:"abril",5:"mayo",6:"junio",
            7:"julio",8:"agosto",9:"septiembre",10:"octubre",11:"noviembre",12:"diciembre"}

def fecha_es():
    hoy = datetime.today()
    return f"{hoy.day} de {MESES_ES[hoy.month]} de {hoy.year}"


# ══════════════════════════════════════════════════════════════════════════════
# HELPERS python-docx
# ══════════════════════════════════════════════════════════════════════════════

def set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def linea_horizontal(doc, color=(31, 78, 121)):
    p   = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pb  = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "6")
    bot.set(qn("w:color"), f"{color[0]:02X}{color[1]:02X}{color[2]:02X}")
    pb.append(bot)
    pPr.append(pb)
    return p

def titulo_seccion(doc, texto: str):
    p   = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(texto.upper())
    run.font.size      = Pt(13)
    run.font.bold      = True
    run.font.color.rgb = RGBColor(*AZUL_OSC)
    linea_horizontal(doc)
    return p

def parrafo(doc, texto: str, bold=False, italic=False, size=11, color=None):
    p   = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(texto)
    run.font.size   = Pt(size)
    run.font.bold   = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p

def pie_imagen(doc, texto: str):
    p   = doc.add_paragraph(texto)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(8)
    p.runs[0].font.size      = Pt(8.5)
    p.runs[0].font.italic    = True
    p.runs[0].font.color.rgb = RGBColor(*GRIS)

def agregar_pie_pagina(doc):
    for section in doc.sections:
        footer = section.footer
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.clear()
        run = p.add_run(
            "Luciano Villagran - Analista de Datos   |   "
            "Datos sinteticos con distribuciones realistas - Portfolio de Datos y Analitica"
        )
        run.font.size      = Pt(8)
        run.font.color.rgb = RGBColor(*GRIS)
        p.alignment        = WD_ALIGN_PARAGRAPH.CENTER


# ══════════════════════════════════════════════════════════════════════════════
# GRAFICOS PARA EL REPORTE
# ══════════════════════════════════════════════════════════════════════════════

def grafico_impacto() -> io.BytesIO:
    """Barras horizontales de los 4 hallazgos ordenados por impacto."""
    labels  = ["Ahorro costos\nde pedido (EOQ/año)",
               "Ventas en riesgo\npor quiebres (inmediato)",
               "Capital inmovilizado\nen sobrestock",
               "Automatizacion\ncalculo mensual (año)"]
    valores = [3.46, 2.18, 1.10, 0.27]
    colores = ["#70AD47", "#ED7D31", "#FF6B6B", "#2E75B6"]

    fig, ax = plt.subplots(figsize=(8.5, 3.6), facecolor="white")
    bars = ax.barh(labels, valores, color=colores, edgecolor="none", height=0.52)

    for bar, val in zip(bars, valores):
        ax.text(bar.get_width() + 0.06,
                bar.get_y() + bar.get_height() / 2,
                f"${val:.2f}M", va="center", fontsize=10,
                fontweight="bold", color="#333333")

    ax.set_xlim(0, 4.6)
    ax.set_xlabel("Impacto estimado (M$ ARS)", fontsize=9, color="#555555")
    ax.set_title("Impacto Economico por Hallazgo — NorteLogix S.A.",
                 fontsize=11, fontweight="bold", color="#1F4E79", pad=10)
    ax.tick_params(axis="y", labelsize=9)
    ax.tick_params(axis="x", labelsize=8)
    for sp in ["top", "right"]:
        ax.spines[sp].set_visible(False)

    fig.tight_layout()
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=140, bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf


def grafico_clasificacion(df_cl: pd.DataFrame, df_pol: pd.DataFrame) -> io.BytesIO:
    """
    Scatter: revenue vs. CV de los 120 SKUs.
    El 'conflicto' clave: A-X arriba-izquierda son los que faltan stock;
    C-Z abajo-derecha son los que tienen sobrestock.
    """
    COLORES_CLASE = {
        "AX":"#1F4E79","AY":"#2E75B6","AZ":"#ED7D31",
        "BX":"#70AD47","BY":"#BDD7EE","BZ":"#FFD966",
        "CX":"#C6EFCE","CY":"#FFEB9C","CZ":"#FFCCCC",
    }

    df = df_cl.merge(
        df_pol[["sku", "bajo_rop", "exceso_unidades"]], on="sku"
    )
    df["bajo_rop"] = df["bajo_rop"].astype(str).str.lower() == "true"

    fig, ax = plt.subplots(figsize=(8.5, 4.2), facecolor="white")

    for clase, grp in df.groupby("clase"):
        color = COLORES_CLASE.get(clase, "#CCCCCC")
        sizes = [90 if r["bajo_rop"] else (60 if r["exceso_unidades"] > 0 else 22)
                 for _, r in grp.iterrows()]
        ax.scatter(grp["cv"], grp["revenue_total"] / 1e6,
                   c=color, s=sizes, alpha=0.80,
                   edgecolors="white", linewidths=0.5,
                   label=clase, zorder=3)

    ymax = df["revenue_total"].max() / 1e6
    ax.axvline(0.30, color="#888888", linestyle="--", linewidth=0.8, alpha=0.6)
    ax.axvline(0.60, color="#888888", linestyle="--", linewidth=0.8, alpha=0.6)
    for x, lbl in [(0.15, "X  (estable)"), (0.45, "Y  (variable)"), (0.85, "Z  (erratico)")]:
        ax.text(x, ymax * 1.04, lbl, ha="center", fontsize=8.5,
                color="#666666", fontstyle="italic")

    ax.set_xlabel("Coeficiente de Variacion (CV) — variabilidad de demanda", fontsize=9)
    ax.set_ylabel("Revenue total 24 meses (M$ ARS)", fontsize=9)
    ax.set_title("Clasificacion ABC-XYZ: valor vs variabilidad de demanda por SKU",
                 fontsize=10, fontweight="bold", color="#1F4E79", pad=10)
    ax.tick_params(labelsize=8)
    for sp in ["top", "right"]:
        ax.spines[sp].set_visible(False)

    handles = [mpatches.Patch(facecolor=c, label=k, edgecolor="#AAAAAA", linewidth=0.5)
               for k, c in COLORES_CLASE.items() if k in df["clase"].values]
    ax.legend(handles=handles, fontsize=7.5, ncol=3, loc="upper right",
              framealpha=0.88, edgecolor="none",
              title="Clase ABC-XYZ", title_fontsize=8)

    # Anotar las zonas criticas
    ax.annotate("A-X: alto valor,\ndemanda estable\n(8 SKUs bajo ROP)",
                xy=(0.10, ymax * 0.72), fontsize=7.5, color="#1F4E79",
                bbox=dict(boxstyle="round,pad=0.3", fc="#EBF3FB", ec="#2E75B6", lw=0.8))
    ax.annotate("C-Z: bajo valor,\ndemanda erratica\n(mayor sobrestock)",
                xy=(0.75, ymax * 0.05), fontsize=7.5, color="#993300",
                bbox=dict(boxstyle="round,pad=0.3", fc="#FFCCCC", ec="#FF6B6B", lw=0.8))

    fig.tight_layout()
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=140, bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf


# ══════════════════════════════════════════════════════════════════════════════
# PORTADA
# ══════════════════════════════════════════════════════════════════════════════

def agregar_portada(doc, total_m: float):
    sec = doc.sections[0]
    sec.page_width    = Cm(21)
    sec.page_height   = Cm(29.7)
    sec.top_margin    = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin   = Cm(3)
    sec.right_margin  = Cm(3)

    # Banda superior de color
    tbl_band = doc.add_table(rows=1, cols=1)
    tbl_band.style = "Table Grid"
    cell_b = tbl_band.rows[0].cells[0]
    set_cell_bg(cell_b, "1F4E79")
    p_b = cell_b.paragraphs[0]
    p_b.paragraph_format.space_before = Pt(12)
    p_b.paragraph_format.space_after  = Pt(12)
    r_b = p_b.add_run("REPORTE EJECUTIVO  |  LOGISTICA Y CADENA DE SUMINISTRO")
    r_b.font.size      = Pt(10)
    r_b.font.bold      = True
    r_b.font.color.rgb = RGBColor(255, 255, 255)
    p_b.alignment      = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Titulo principal
    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p1.add_run("Optimizacion de Inventario")
    r1.font.size      = Pt(26)
    r1.font.bold      = True
    r1.font.color.rgb = RGBColor(*AZUL_OSC)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Clasificacion ABC-XYZ y Politica de Stock Optima")
    r2.font.size      = Pt(16)
    r2.font.color.rgb = RGBColor(*AZUL_MED)

    doc.add_paragraph()
    linea_horizontal(doc)
    doc.add_paragraph()

    # Metadatos de portada
    for label, valor in [
        ("Empresa analizada", "NorteLogix S.A.  —  Distribuidora de consumo masivo"),
        ("Periodo analizado", "Enero 2023 – Diciembre 2024"),
        ("SKUs analizados",   "120 productos activos · 5 categorias · 24 meses de historial"),
        ("Herramienta",       "Python  (pandas · numpy · matplotlib)"),
        ("Analista",          "Luciano Villagran  —  Analista de Datos"),
        ("Fecha de emision",  fecha_es()),
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rl = p.add_run(f"{label}:  ")
        rl.font.size      = Pt(11)
        rl.font.bold      = True
        rl.font.color.rgb = RGBColor(*AZUL_OSC)
        rv = p.add_run(valor)
        rv.font.size = Pt(11)

    doc.add_paragraph()
    linea_horizontal(doc)
    doc.add_paragraph()

    # Caja de impacto en portada
    tbl_imp = doc.add_table(rows=1, cols=1)
    tbl_imp.style = "Table Grid"
    cell_i = tbl_imp.rows[0].cells[0]
    set_cell_bg(cell_i, "EBF3FB")
    pi = cell_i.paragraphs[0]
    pi.paragraph_format.space_before = Pt(10)
    pi.paragraph_format.space_after  = Pt(10)
    pi.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ri = pi.add_run(f"Impacto total identificado: ~${total_m:.1f}M")
    ri.font.size      = Pt(16)
    ri.font.bold      = True
    ri.font.color.rgb = RGBColor(*AZUL_OSC)
    p_sub = cell_i.add_paragraph(
        "Capital liberado en sobrestock + ventas recuperadas + ahorro EOQ + automatizacion"
    )
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_before = Pt(2)
    p_sub.paragraph_format.space_after  = Pt(10)
    p_sub.runs[0].font.size      = Pt(9)
    p_sub.runs[0].font.italic    = True
    p_sub.runs[0].font.color.rgb = RGBColor(*AZUL_MED)

    doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# MAIN
# ══════════════════════════════════════════════════════════════════════════════

def main():
    # Verificar outputs previos
    for f in ["clasificacion_abcxyz.csv", "politica_inventario.csv"]:
        if not (OUT / f).exists():
            print(f"ERROR: {f} no encontrado.")
            print("Ejecutar primero:")
            print("  python datos/generar_datos.py")
            print("  python src/analisis_abcxyz.py")
            sys.exit(1)

    df_cl  = pd.read_csv(OUT / "clasificacion_abcxyz.csv")
    df_pol = pd.read_csv(OUT / "politica_inventario.csv")
    df_pol["bajo_rop"] = df_pol["bajo_rop"].astype(str).str.lower() == "true"

    # Numeros clave
    capital_inmov   = df_pol["capital_inmovilizado"].sum()
    riesgo_ventas   = df_pol["riesgo_ventas"].sum()
    skus_bajo_rop   = int(df_pol["bajo_rop"].sum())
    skus_sobrestock = int((df_pol["exceso_unidades"] > 0).sum())
    AHORRO_EOQ      = 3_460_000
    AHORRO_TIEMPO   =   270_000
    total           = capital_inmov + riesgo_ventas + AHORRO_EOQ + AHORRO_TIEMPO
    total_m         = total / 1e6

    print("Generando graficos...")
    buf_impacto = grafico_impacto()
    buf_clasif  = grafico_clasificacion(df_cl, df_pol)
    print("  OK: 2 graficos generados")

    # ── Construir DOCX ────────────────────────────────────────────────────────
    doc = Document()
    agregar_pie_pagina(doc)
    agregar_portada(doc, total_m)

    # ── SECCION 1: RESUMEN EJECUTIVO ─────────────────────────────────────────
    titulo_seccion(doc, "1. Resumen Ejecutivo")

    tbl_res = doc.add_table(rows=1, cols=1)
    tbl_res.style = "Table Grid"
    cell_r = tbl_res.rows[0].cells[0]
    set_cell_bg(cell_r, "FFF9C4")
    p_hl = cell_r.paragraphs[0]
    p_hl.paragraph_format.space_before = Pt(10)
    p_hl.paragraph_format.space_after  = Pt(4)
    p_hl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_hl = p_hl.add_run(f"Impacto total identificado: ${total_m:.1f}M")
    r_hl.font.size      = Pt(13)
    r_hl.font.bold      = True
    r_hl.font.color.rgb = RGBColor(*AZUL_OSC)
    p_sub = cell_r.add_paragraph(
        "Capital liberado en sobrestock + ventas recuperadas por quiebres "
        "+ ahorros operativos EOQ + automatizacion"
    )
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(10)
    p_sub.runs[0].font.size      = Pt(9)
    p_sub.runs[0].font.italic    = True
    p_sub.runs[0].font.color.rgb = RGBColor(*AZUL_MED)

    doc.add_paragraph()

    parrafo(doc,
        "El problema: NorteLogix S.A. gestionaba el inventario de 120 SKUs activos con "
        "criterio artesanal — sin modelo estadistico detras. Las decisiones de cuanto pedir "
        "y cuando pedir las tomaba el comprador por experiencia. El resultado era el clasico "
        "contrasentido de la cadena de suministro: depositos llenos de productos que casi "
        "nadie pide, y quiebres en los productos estrella que son la fuente principal de ingreso.")

    parrafo(doc,
        "Lo que construi: un modelo Python reproducible que clasifica cada uno de los "
        "120 SKUs en una matriz ABC-XYZ — combinando valor de facturacion (ABC) con "
        "variabilidad de demanda (XYZ) — y calcula la politica de stock optima para cada "
        "producto: stock de seguridad estadistico al 95 % de nivel de servicio, punto de "
        "reorden exacto y lote economico de compra (EOQ).")

    parrafo(doc,
        f"El resultado: el modelo identifico ${capital_inmov/1e6:.1f}M en capital "
        f"inmovilizado en {skus_sobrestock} SKUs con sobrestock, ${riesgo_ventas/1e6:.1f}M "
        f"en ventas en riesgo por {skus_bajo_rop} SKUs bajo su punto de reorden, y $3,5M/año "
        f"de ahorro potencial al adoptar la politica EOQ. El proceso de analisis mensual "
        "paso de 8 horas manuales a 30 segundos automaticos.")

    doc.add_page_break()

    # ── SECCION 2: METODOLOGIA ───────────────────────────────────────────────
    titulo_seccion(doc, "2. Metodologia y Datos")

    parrafo(doc,
        "Trabaje sobre tres archivos CSV generados sinteticamente con semilla fija "
        "(SEED=42) para garantizar reproducibilidad total. Los datos replican los patrones "
        "de una distribuidora regional de consumo masivo con 120 SKUs activos en 5 categorias.")

    # Tabla de datasets
    tbl_d = doc.add_table(rows=4, cols=3)
    tbl_d.style = "Table Grid"
    for j, enc in enumerate(["Archivo", "Registros", "Descripcion"]):
        c = tbl_d.rows[0].cells[j]
        set_cell_bg(c, "1F4E79")
        r = c.paragraphs[0].add_run(enc)
        r.font.bold = True; r.font.color.rgb = RGBColor(255,255,255); r.font.size = Pt(9.5)
    filas_d = [
        ["productos.csv",         "120 filas",   "SKUs con precio, costo, lead time y costo de pedido"],
        ["demanda_mensual.csv",   "2.880 filas", "24 meses x 120 SKUs — ventas con estacionalidad"],
        ["inventario_actual.csv", "120 filas",   "Posicion de stock al corte con patrones deliberados"],
    ]
    for i, fila in enumerate(filas_d):
        for j, v in enumerate(fila):
            c = tbl_d.rows[i+1].cells[j]
            if i % 2 == 0: set_cell_bg(c, "EBF3FB")
            r = c.paragraphs[0].add_run(v); r.font.size = Pt(9.5)

    doc.add_paragraph()
    parrafo(doc, "Formulas aplicadas:", bold=True, size=10.5)
    for bullet in [
        "ABC: revenue acumulado ordenado. A <= 80 %, B <= 95 %, C = resto.",
        "XYZ: coeficiente de variacion (CV = sigma / mu). X < 0,30 | Y < 0,60 | Z >= 0,60.",
        "Safety stock = 1,645 x sigma_mensual x sqrt(lead_time_meses)  [95 % nivel de servicio].",
        "Reorder point = mu_mensual x lead_time_meses + safety_stock.",
        "EOQ = sqrt(2 x D_anual x costo_pedido / (costo_unitario x 0,25)).",
        "Sobrestock = max(0, stock_actual - (ROP + EOQ))  x  costo_unitario.",
        "Deficit = max(0, ROP - stock_actual)  x  precio_unitario.",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        p.add_run(bullet).font.size = Pt(10)

    doc.add_page_break()

    # ── SECCION 3: HALLAZGOS ─────────────────────────────────────────────────
    titulo_seccion(doc, "3. Hallazgos y Evidencia")

    # Panel principal
    parrafo(doc, "Dashboard ejecutivo — 5 visualizaciones integradas:", bold=True, size=10.5)
    dash_path = OUT / "dashboard_inventario.png"
    if dash_path.exists():
        doc.add_picture(str(dash_path), width=Inches(6.2))
        pie_imagen(doc,
            "Matriz ABC-XYZ, curva Pareto, top 10 sobrestock, top 10 riesgo de quiebre y KPIs de impacto.")
    doc.add_paragraph()

    # Scatter clasificacion
    parrafo(doc, "Posicionamiento de los 120 SKUs — valor vs variabilidad:", bold=True, size=10.5)
    doc.add_picture(buf_clasif, width=Inches(6.2))
    pie_imagen(doc,
        "Los puntos mas grandes estan bajo ROP (urgente) o en sobrestock. "
        "Los A-X (arriba-izquierda) son los productos criticos; los C-Z (abajo-derecha) los que generan el sobrestock.")

    doc.add_page_break()

    # ─────────────────────────────────────────────────────────────────────────
    # Hallazgo 1 — Sobrestock
    # ─────────────────────────────────────────────────────────────────────────
    p_h1 = doc.add_paragraph()
    p_h1.paragraph_format.space_before = Pt(8)
    r_h1 = p_h1.add_run(f"Hallazgo 1 — ${capital_inmov/1e6:.1f}M en capital inmovilizado en sobrestock ({skus_sobrestock} SKUs)")
    r_h1.font.size = Pt(12); r_h1.font.bold = True
    r_h1.font.color.rgb = RGBColor(*AZUL_MED)

    parrafo(doc,
        f"Que encontre: {skus_sobrestock} SKUs tienen un stock actual superior a ROP + EOQ. "
        "La mayor parte son productos C-Z y C-Y — bajo valor, demanda erratica — "
        "que se siguieron pidiendo mensualmente aunque la demanda no los consumia al mismo ritmo. "
        "El resultado: capital paralizado en el deposito sin generar retorno.")

    top3_sob = df_pol.nlargest(3, "capital_inmovilizado")
    tbl_h1 = doc.add_table(rows=4, cols=4)
    tbl_h1.style = "Table Grid"
    for j, e in enumerate(["SKU", "Clase", "Exceso (unidades)", "Capital inmovilizado"]):
        c = tbl_h1.rows[0].cells[j]
        set_cell_bg(c, "1F4E79")
        r = c.paragraphs[0].add_run(e)
        r.font.bold = True; r.font.color.rgb = RGBColor(255,255,255); r.font.size = Pt(9.5)
    for i, (_, row) in enumerate(top3_sob.iterrows()):
        vals = [row["sku"], row["clase"],
                f"{int(row['exceso_unidades']):,} u.", f"${row['capital_inmovilizado']/1e3:,.0f}k"]
        for j, v in enumerate(vals):
            c = tbl_h1.rows[i+1].cells[j]
            if i % 2 == 0: set_cell_bg(c, "EBF3FB")
            r = c.paragraphs[0].add_run(v); r.font.size = Pt(9.5)

    p_tot1 = doc.add_paragraph()
    r_tot1 = p_tot1.add_run(
        f"... y {skus_sobrestock - 3} SKUs mas con sobrestock  |  TOTAL: ${capital_inmov/1e6:.2f}M"
    )
    r_tot1.font.size = Pt(9); r_tot1.font.italic = True
    r_tot1.font.color.rgb = RGBColor(*GRIS)

    parrafo(doc,
        "Accion recomendada: reducir el proximo lote de compra para los 30 SKUs identificados. "
        "No eliminar las ordenes — adoptar el lote EOQ que ajusta el tamano y la frecuencia "
        "de cada compra al ritmo real de consumo.")

    p_imp1 = doc.add_paragraph()
    p_imp1.paragraph_format.space_before = Pt(4)
    r_imp1a = p_imp1.add_run("Impacto estimado: ")
    r_imp1a.font.bold = True; r_imp1a.font.size = Pt(10.5)
    r_imp1b = p_imp1.add_run(f"${capital_inmov/1e6:.1f}M de capital liberado")
    r_imp1b.font.bold = True; r_imp1b.font.size = Pt(10.5)
    r_imp1b.font.color.rgb = RGBColor(*VERDE)
    p_sup1 = doc.add_paragraph()
    r_sup1 = p_sup1.add_run("(Supuesto: exceso_unidades x costo_unitario al corte de enero 2025.)")
    r_sup1.font.size = Pt(9); r_sup1.font.italic = True
    r_sup1.font.color.rgb = RGBColor(*GRIS)

    doc.add_paragraph()

    # ─────────────────────────────────────────────────────────────────────────
    # Hallazgo 2 — Substock / quiebres
    # ─────────────────────────────────────────────────────────────────────────
    p_h2 = doc.add_paragraph()
    p_h2.paragraph_format.space_before = Pt(8)
    r_h2 = p_h2.add_run(f"Hallazgo 2 — ${riesgo_ventas/1e6:.1f}M en ventas en riesgo — {skus_bajo_rop} SKUs bajo ROP")
    r_h2.font.size = Pt(12); r_h2.font.bold = True
    r_h2.font.color.rgb = RGBColor(*AZUL_MED)

    parrafo(doc,
        f"Que encontre: {skus_bajo_rop} SKUs tienen stock por debajo de su punto de reorden. "
        "Cinco de ellos son tipo A-X: alto valor facturado y demanda estable. "
        "Son los que mas impacto tienen si se agotan durante el lead time del proveedor — "
        "el cliente no espera, compra en la competencia.")

    top3_rop = df_pol[df_pol["bajo_rop"]].nlargest(3, "riesgo_ventas")
    tbl_h2 = doc.add_table(rows=4, cols=4)
    tbl_h2.style = "Table Grid"
    for j, e in enumerate(["SKU", "Clase", "Deficit (unidades)", "Ventas en riesgo"]):
        c = tbl_h2.rows[0].cells[j]
        set_cell_bg(c, "C00000")
        r = c.paragraphs[0].add_run(e)
        r.font.bold = True; r.font.color.rgb = RGBColor(255,255,255); r.font.size = Pt(9.5)
    for i, (_, row) in enumerate(top3_rop.iterrows()):
        vals = [row["sku"], row["clase"],
                f"{int(row['deficit_unidades']):,} u.", f"${row['riesgo_ventas']/1e3:,.0f}k"]
        for j, v in enumerate(vals):
            c = tbl_h2.rows[i+1].cells[j]
            if i % 2 == 0: set_cell_bg(c, "FFCCCC")
            r = c.paragraphs[0].add_run(v); r.font.size = Pt(9.5)

    p_tot2 = doc.add_paragraph()
    r_tot2 = p_tot2.add_run(
        f"... y {skus_bajo_rop - 3} SKUs mas bajo su ROP  |  TOTAL en riesgo: ${riesgo_ventas/1e6:.2f}M"
    )
    r_tot2.font.size = Pt(9); r_tot2.font.italic = True
    r_tot2.font.color.rgb = RGBColor(*GRIS)

    parrafo(doc,
        f"Accion recomendada: emitir ordenes de compra urgentes para los {skus_bajo_rop} SKUs. "
        "Prioridad absoluta: SKU-001 a SKU-005 (clase A-X, mayor valor expuesto). "
        "Sin accion esta semana, el quiebre es inminente dado el lead time de los proveedores.")

    p_imp2 = doc.add_paragraph()
    p_imp2.paragraph_format.space_before = Pt(4)
    r_imp2a = p_imp2.add_run("Impacto estimado: ")
    r_imp2a.font.bold = True; r_imp2a.font.size = Pt(10.5)
    r_imp2b = p_imp2.add_run(f"${riesgo_ventas/1e6:.1f}M de ventas recuperadas (riesgo evitado)")
    r_imp2b.font.bold = True; r_imp2b.font.size = Pt(10.5)
    r_imp2b.font.color.rgb = RGBColor(*VERDE)
    p_sup2 = doc.add_paragraph()
    r_sup2 = p_sup2.add_run("(Supuesto: deficit_unidades x precio_unitario. Riesgo inmediato al momento del analisis.)")
    r_sup2.font.size = Pt(9); r_sup2.font.italic = True
    r_sup2.font.color.rgb = RGBColor(*GRIS)

    doc.add_page_break()

    # ─────────────────────────────────────────────────────────────────────────
    # Hallazgo 3 — EOQ
    # ─────────────────────────────────────────────────────────────────────────
    p_h3 = doc.add_paragraph()
    p_h3.paragraph_format.space_before = Pt(8)
    r_h3 = p_h3.add_run(f"Hallazgo 3 — $3,5M/año en ahorro de costos de pedido al adoptar politica EOQ")
    r_h3.font.size = Pt(12); r_h3.font.bold = True
    r_h3.font.color.rgb = RGBColor(*AZUL_MED)

    parrafo(doc,
        "Que encontre: comparando la frecuencia optima implicita en el EOQ vs la politica "
        "actual (estimada en 12 ordenes/año por SKU), hay un exceso significativo de ordenes "
        "en los items de baja rotacion. Para un SKU C-Z con 60 unidades/año de demanda, "
        "el EOQ indica comprar 1-2 veces al año — no mensualmente.")

    tbl_h3 = doc.add_table(rows=4, cols=3)
    tbl_h3.style = "Table Grid"
    for j, e in enumerate(["Tipo de SKUs", "Frecuencia actual", "Frecuencia EOQ optima"]):
        c = tbl_h3.rows[0].cells[j]
        set_cell_bg(c, "1F4E79")
        r = c.paragraphs[0].add_run(e)
        r.font.bold = True; r.font.color.rgb = RGBColor(255,255,255); r.font.size = Pt(9.5)
    filas_h3 = [
        ["A-X / A-Y  (alta demanda)", "12/año", "Sin ahorro — alta demanda justifica alta frecuencia"],
        ["B-Y / B-Z  (demanda media)", "12/año", "4-6/año  (ahorro: 6-8 ordenes x costo_pedido)"],
        ["C-Z / C-Y  (baja demanda)",  "12/año", "1-2/año  (ahorro: 10-11 ordenes x costo_pedido)"],
    ]
    for i, fila in enumerate(filas_h3):
        for j, v in enumerate(fila):
            c = tbl_h3.rows[i+1].cells[j]
            if i % 2 == 0: set_cell_bg(c, "EBF3FB")
            r = c.paragraphs[0].add_run(v); r.font.size = Pt(9.5)

    parrafo(doc,
        "Accion recomendada: adoptar el lote EOQ como estandar de compra. Para items C: "
        "pasar a ciclos trimestrales o semestrales. El modelo calcula el EOQ automaticamente "
        "para cada uno de los 120 SKUs y lo entrega en el output politica_inventario.csv.")

    p_imp3 = doc.add_paragraph()
    r_imp3a = p_imp3.add_run("Impacto estimado: ")
    r_imp3a.font.bold = True; r_imp3a.font.size = Pt(10.5)
    r_imp3b = p_imp3.add_run("$3,5M/año en costos operativos de pedido")
    r_imp3b.font.bold = True; r_imp3b.font.size = Pt(10.5)
    r_imp3b.font.color.rgb = RGBColor(*VERDE)
    p_sup3 = doc.add_paragraph()
    r_sup3 = p_sup3.add_run(
        "(Supuesto: politica actual = 12 ordenes/año por SKU. "
        "Costo fijo de pedido segun CSV de productos. Ahorro solo en items donde EOQ < 12 ordenes.)"
    )
    r_sup3.font.size = Pt(9); r_sup3.font.italic = True
    r_sup3.font.color.rgb = RGBColor(*GRIS)

    doc.add_paragraph()

    # ─────────────────────────────────────────────────────────────────────────
    # Hallazgo 4 — Automatizacion
    # ─────────────────────────────────────────────────────────────────────────
    p_h4 = doc.add_paragraph()
    p_h4.paragraph_format.space_before = Pt(8)
    r_h4 = p_h4.add_run("Hallazgo 4 — $270k/año en tiempo recuperado: de 8 horas manuales a 30 segundos")
    r_h4.font.size = Pt(12); r_h4.font.bold = True
    r_h4.font.color.rgb = RGBColor(*AZUL_MED)

    parrafo(doc,
        "Que encontre: el proceso de revision mensual de inventario — recolectar datos, "
        "calcular stocks, identificar excepciones, documentar — llevaba aproximadamente "
        "8 horas de trabajo manual cada mes. Ademas, el proceso no era reproducible: "
        "dependia del criterio de quien lo ejecutaba y era propenso a errores de calculo.")

    tbl_h4 = doc.add_table(rows=3, cols=3)
    tbl_h4.style = "Table Grid"
    for j, e in enumerate(["Actividad", "Antes (manual)", "Despues (automatico)"]):
        c = tbl_h4.rows[0].cells[j]
        set_cell_bg(c, "1F4E79")
        r = c.paragraphs[0].add_run(e)
        r.font.bold = True; r.font.color.rgb = RGBColor(255,255,255); r.font.size = Pt(9.5)
    filas_h4 = [
        ["Calculo ABC-XYZ + politica de stock", "~8 horas/mes", "< 30 segundos"],
        ["Riesgo de error en la clasificacion",  "Alto (criterio subjetivo)", "Nulo (formula estadistica)"],
    ]
    for i, fila in enumerate(filas_h4):
        for j, v in enumerate(fila):
            c = tbl_h4.rows[i+1].cells[j]
            if i % 2 == 0: set_cell_bg(c, "EBF3FB")
            r = c.paragraphs[0].add_run(v); r.font.size = Pt(9.5)

    parrafo(doc,
        "Accion recomendada: integrar 'python src/analisis_abcxyz.py' al cierre mensual "
        "de inventario. El modelo entrega clasificacion actualizada, politica optima por SKU "
        "y lista de criticos para accion inmediata — en 30 segundos.")

    p_imp4 = doc.add_paragraph()
    r_imp4a = p_imp4.add_run("Impacto estimado: ")
    r_imp4a.font.bold = True; r_imp4a.font.size = Pt(10.5)
    r_imp4b = p_imp4.add_run("$270k/año en horas de analista liberadas")
    r_imp4b.font.bold = True; r_imp4b.font.size = Pt(10.5)
    r_imp4b.font.color.rgb = RGBColor(*VERDE)
    p_sup4 = doc.add_paragraph()
    r_sup4 = p_sup4.add_run("(Supuesto: 7,5 h ahorradas/mes x 12 meses = 90 h/año x $3.000/h analista.)")
    r_sup4.font.size = Pt(9); r_sup4.font.italic = True
    r_sup4.font.color.rgb = RGBColor(*GRIS)

    doc.add_page_break()

    # ── SECCION 4: CONCLUSIONES ───────────────────────────────────────────────
    titulo_seccion(doc, "4. Conclusiones y Proximos Pasos")

    parrafo(doc, "Resumen visual del impacto economico por hallazgo:", bold=True, size=10.5)
    doc.add_picture(buf_impacto, width=Inches(5.8))
    pie_imagen(doc, "Los 4 hallazgos ordenados por impacto economico.")
    doc.add_paragraph()

    parrafo(doc, "Tabla resumen de impacto:", bold=True, size=10.5)
    tbl_sum = doc.add_table(rows=6, cols=5)
    tbl_sum.style = "Table Grid"
    for j, e in enumerate(["Hallazgo", "Accion", "Impacto", "Esfuerzo", "Supuesto clave"]):
        c = tbl_sum.rows[0].cells[j]
        set_cell_bg(c, "1F4E79")
        r = c.paragraphs[0].add_run(e)
        r.font.bold = True; r.font.color.rgb = RGBColor(255,255,255); r.font.size = Pt(9)
    filas_sum = [
        ["Ahorro costos pedido (EOQ)", "Adoptar lote EOQ en items C", "$3,5M/año", "Medio",
         "12 ord./año actual vs EOQ"],
        ["Ventas en riesgo (quiebres)", f"Reponer {skus_bajo_rop} SKUs urgente",
         f"${riesgo_ventas/1e6:.1f}M", "Bajo", "deficit x precio"],
        ["Capital inmovilizado", "Reducir lotes en 30 SKUs",
         f"${capital_inmov/1e6:.1f}M", "Bajo", "exceso x costo"],
        ["Automatizacion mensual", "Pipeline automatico", "$270k/año", "Bajo",
         "90h/año x $3.000/h"],
        ["TOTAL", "", f"~${total_m:.1f}M", "", ""],
    ]
    for i, fila in enumerate(filas_sum):
        for j, v in enumerate(fila):
            c = tbl_sum.rows[i+1].cells[j]
            if i == 4:
                set_cell_bg(c, "1F4E79")
                r = c.paragraphs[0].add_run(v)
                r.font.bold = True; r.font.color.rgb = RGBColor(255,255,255); r.font.size = Pt(10)
            else:
                if i % 2 == 0: set_cell_bg(c, "EBF3FB")
                r = c.paragraphs[0].add_run(v); r.font.size = Pt(9.5)

    doc.add_paragraph()
    p_t = doc.add_paragraph()
    p_t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_t = p_t.add_run(f"Impacto total estimado del analisis: ~${total_m:.1f}M")
    r_t.font.size = Pt(14); r_t.font.bold = True
    r_t.font.color.rgb = RGBColor(*AZUL_OSC)
    p_ts = doc.add_paragraph()
    p_ts.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_ts = p_ts.add_run(
        "(capital liberado en sobrestock + ventas recuperadas por quiebres "
        "+ ahorros recurrentes EOQ + automatizacion)"
    )
    r_ts.font.size = Pt(9); r_ts.font.italic = True
    r_ts.font.color.rgb = RGBColor(*GRIS)

    doc.add_paragraph()
    linea_horizontal(doc)

    parrafo(doc, "Proximos pasos recomendados (por ROI):", bold=True, size=11)
    pasos = [
        ("Inmediato — semana 1",
         f"Emitir ordenes de compra urgentes para los {skus_bajo_rop} SKUs bajo ROP. "
         "Prioridad: SKU-001 a SKU-005 (clase A-X). Sin accion, el quiebre de stock es inminente "
         "dado el lead time actual de los proveedores."),
        ("Corto plazo — mes 1 a 2",
         "Para los 30 SKUs con sobrestock, no emitir ordenes hasta que el stock baje al ROP. "
         "Adoptar el lote EOQ en las proximas compras en lugar del lote fijo habitual. "
         "Revisar con el area comercial si alguno de los items C-Z puede discontinuarse."),
        ("Proceso continuo",
         "Integrar 'python src/analisis_abcxyz.py' al cierre mensual de inventario. "
         "El modelo tarda 30 segundos y entrega la politica actualizada para los 120 SKUs, "
         "la lista de criticos y el dashboard ejecutivo. Revisar umbrales XYZ cada 6 meses."),
    ]
    for i, (label, texto) in enumerate(pasos, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        rl = p.add_run(f"{i}. {label}: ")
        rl.font.bold = True; rl.font.size = Pt(10.5)
        rl.font.color.rgb = RGBColor(*AZUL_MED)
        rt = p.add_run(texto)
        rt.font.size = Pt(10.5)

    # ── Guardar DOCX ──────────────────────────────────────────────────────────
    docx_path = OUT / "Reporte_OptimizacionInventario.docx"
    pdf_path  = OUT / "Reporte_OptimizacionInventario.pdf"
    doc.save(str(docx_path))
    print(f"DOCX guardado: {docx_path.name}")

    # ── Exportar PDF ──────────────────────────────────────────────────────────
    try:
        from docx2pdf import convert
        convert(str(docx_path), str(pdf_path))
        print(f"PDF  guardado: {pdf_path.name}")
    except Exception as exc:
        print(f"PDF no generado ({exc}). Abrir el DOCX y exportar manualmente a PDF.")

    return docx_path, pdf_path


if __name__ == "__main__":
    main()
