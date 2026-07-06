"""
generar_reporte.py
Genera el reporte ejecutivo del análisis de Rentabilidad Comercial — TiendaNova.
Output: output/Reporte_RentabilidadComercial.docx + .pdf

Ejecutar: python generar_reporte.py
Requiere: python-docx, docx2pdf, matplotlib
"""

import os
import io
from pathlib import Path
from datetime import date

# ── python-docx ──────────────────────────────────────────────────────────────
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── matplotlib ───────────────────────────────────────────────────────────────
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches

BASE   = Path(__file__).resolve().parent
OUT    = BASE / "output"
OUT.mkdir(exist_ok=True)


def _ar(x, dec=0):
    """Número en formato argentino: miles con punto, decimales con coma."""
    return f"{x:,.{dec}f}".replace(",", "§").replace(".", ",").replace("§", ".")
RUTA_DOCX = OUT / "Reporte_RentabilidadComercial.docx"
RUTA_PDF  = OUT / "Reporte_RentabilidadComercial.pdf"
IMG_EXIST = BASE / "output" / "impacto_hallazgos.png"

# ══════════════════════════════════════════════════════════════════════════════
# COLORES
# ══════════════════════════════════════════════════════════════════════════════
C_AZUL_OSC = RGBColor(0x1F, 0x4E, 0x79)
C_AZUL_MED = RGBColor(0x2E, 0x75, 0xB6)
C_NARANJA  = RGBColor(0xED, 0x7D, 0x31)
C_VERDE    = RGBColor(0x70, 0xAD, 0x47)
C_GRIS     = RGBColor(0x59, 0x59, 0x59)
C_BLANCO   = RGBColor(0xFF, 0xFF, 0xFF)

HEX_AZUL_OSC = "1F4E79"
HEX_AZUL_MED = "2E75B6"
HEX_AZUL_CLA = "BDD7EE"
HEX_VERDE_CL = "C6EFCE"
HEX_ROJO_CL  = "FFCCCC"
HEX_AMAR_CL  = "FFEB9C"
HEX_DIAG     = "FFF2CC"

# ══════════════════════════════════════════════════════════════════════════════
# HELPERS
# ══════════════════════════════════════════════════════════════════════════════
def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def set_para_bg(para, hex_color):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    pPr.append(shd)

def linea_horizontal(doc, color_hex=HEX_AZUL_MED):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)

def titulo_seccion(doc, texto, nivel=1):
    h = doc.add_heading("", level=nivel)
    run = h.add_run(texto)
    run.bold = True
    run.font.color.rgb = C_AZUL_OSC
    run.font.size = Pt(14 if nivel == 1 else 12)
    return h

def parrafo(doc, texto, size=11, bold=False, color=None, align=WD_ALIGN_PARAGRAPH.LEFT, italic=False, espacio_antes=0):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(espacio_antes)
    run = p.add_run(texto)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.color.rgb = color if color else C_GRIS
    return p

def bloque_hallazgo(doc, numero, titulo, que_encontre, accion, impacto, supuesto, color_imp=None):
    """Bloque estándar por hallazgo: título destacado + campos."""
    # Título del hallazgo con fondo azul claro
    p_tit = doc.add_paragraph()
    p_tit.paragraph_format.space_before = Pt(8)
    p_tit.paragraph_format.space_after  = Pt(2)
    set_para_bg(p_tit, HEX_AZUL_CLA)
    run_n = p_tit.add_run(f"  ► Hallazgo {numero} — ")
    run_n.bold = True
    run_n.font.size = Pt(11)
    run_n.font.color.rgb = C_AZUL_OSC
    run_t = p_tit.add_run(titulo)
    run_t.bold = True
    run_t.font.size = Pt(11)
    run_t.font.color.rgb = C_AZUL_OSC

    def campo(etiqueta, valor, color_val=None):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent  = Cm(0.5)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        r1 = p.add_run(f"{etiqueta}: ")
        r1.bold = True
        r1.font.size = Pt(10)
        r1.font.color.rgb = RGBColor(0x3D, 0x3D, 0x3D)
        r2 = p.add_run(valor)
        r2.font.size = Pt(10)
        r2.font.color.rgb = color_val if color_val else RGBColor(0x3D, 0x3D, 0x3D)

    campo("Qué encontré", que_encontre)
    campo("Acción recomendada", accion)
    campo("Impacto estimado", impacto, color_val=C_VERDE if not color_imp else color_imp)
    campo("Supuesto clave", supuesto)

# ══════════════════════════════════════════════════════════════════════════════
# GRÁFICO 1 — Impacto por hallazgo (barras horizontales)
# ══════════════════════════════════════════════════════════════════════════════
def generar_grafico_impacto():
    hallazgos = [
        "Productos bajo\nmargen (repreciar)",
        "Clientes en fuga\n(campaña recuperación)",
        "Brecha sucursal Sur\n(replicar mejores prácticas)",
        "Baja rotación\n(capital liberado)",
    ]
    valores  = [0.91, 0.74, 0.39, 0.006]
    colores  = ["#2E75B6", "#ED7D31", "#70AD47", "#BDD7EE"]

    fig, ax = plt.subplots(figsize=(9, 4.5))
    fig.patch.set_facecolor("white")
    bars = ax.barh(range(len(hallazgos)), valores, color=colores, height=0.55,
                   edgecolor="white", linewidth=0.5, zorder=2)
    ax.set_yticks(range(len(hallazgos)))
    ax.set_yticklabels(list(reversed(hallazgos)) if False else hallazgos, fontsize=9.5)
    ax.set_xlabel("Impacto estimado (millones de $ ARS / año)", fontsize=10)
    ax.set_title("Impacto económico identificado por hallazgo — TiendaNova 2024–2025",
                 fontsize=11, fontweight="bold", color="#1F4E79", pad=10)
    ax.xaxis.set_major_formatter(plt.FuncFormatter(lambda x, _: f"$ {_ar(x, 2)} M"))
    ax.xaxis.grid(True, linestyle="--", alpha=0.35, zorder=0)
    ax.set_axisbelow(True)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    for i, val in enumerate(valores):
        label = f"${_ar(val, 2)} M/año" if val >= 0.05 else f"~${_ar(79, 0)} k (puntual)"
        ax.text(val + 0.01, i, label, va="center", fontsize=9, fontweight="bold",
                color="#1F4E79")
    ax.axvline(0.91 + 0.74 + 0.39, color="#595959", linestyle=":", linewidth=1, alpha=0.6)
    ax.text(0.91 + 0.74 + 0.39 + 0.01, 3.4, "Total\n~$2,0 M/año",
            fontsize=8, color="#595959", va="top")
    plt.tight_layout()
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150, bbox_inches="tight")
    buf.seek(0)
    plt.close(fig)
    return buf

# ══════════════════════════════════════════════════════════════════════════════
# GRÁFICO 2 — Ticket promedio por sucursal
# ══════════════════════════════════════════════════════════════════════════════
def generar_grafico_sucursales():
    sucursales = ["Sur\n(Bariloche)", "Cuyo", "Centro", "Litoral", "Norte"]
    tickets    = [9499, 10882, 11267, 11730, 12713]
    colores    = ["#FF6B6B", "#FFD966", "#70AD47", "#70AD47", "#2E75B6"]

    fig, ax = plt.subplots(figsize=(9, 4))
    fig.patch.set_facecolor("white")
    ax.bar(range(len(sucursales)), tickets, color=colores, edgecolor="white",
           linewidth=0.5, zorder=2)
    ax.set_xticks(range(len(sucursales)))
    ax.set_xticklabels(sucursales, fontsize=9.5)
    ax.set_ylabel("Ticket promedio ($ ARS)", fontsize=10)
    ax.set_title("Ticket promedio por sucursal — TiendaNova 2024–2025\nSucursal Sur: −16 % vs mediana de la red",
                 fontsize=11, fontweight="bold", color="#1F4E79", pad=10)
    ax.yaxis.set_major_formatter(plt.FuncFormatter(lambda x, _: f"${_ar(x, 0)}"))
    ax.yaxis.grid(True, linestyle="--", alpha=0.35, zorder=0)
    ax.set_axisbelow(True)
    ax.axhline(11267, color="#595959", linestyle="--", linewidth=1.1, alpha=0.6)
    ax.text(4.5, 11350, f"Mediana red\n${_ar(11267, 0)}", fontsize=8, color="#595959", ha="right")
    for i, val in enumerate(tickets):
        ax.text(i, val + 100, f"${_ar(val, 0)}", ha="center", fontsize=8.5, fontweight="bold",
                color="#1F4E79")
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.set_ylim(8000, 14000)
    legend_elems = [
        mpatches.Patch(facecolor="#FF6B6B", label="Por debajo de la red"),
        mpatches.Patch(facecolor="#FFD966", label="Cercano a la mediana"),
        mpatches.Patch(facecolor="#70AD47", label="En la mediana"),
        mpatches.Patch(facecolor="#2E75B6", label="Mejor desempeño"),
    ]
    ax.legend(handles=legend_elems, fontsize=8, loc="lower right", framealpha=0.85)
    plt.tight_layout()
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150, bbox_inches="tight")
    buf.seek(0)
    plt.close(fig)
    return buf

# ══════════════════════════════════════════════════════════════════════════════
# CONSTRUIR DOCUMENTO
# ══════════════════════════════════════════════════════════════════════════════
doc = Document()

# Márgenes
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(2.8)
    section.right_margin  = Cm(2.8)

# ── PIE DE PÁGINA ─────────────────────────────────────────────────────────────
section = doc.sections[0]
footer  = section.footer
fp = footer.paragraphs[0]
fp.clear()
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = fp.add_run("Luciano Villagrán — Analista de Datos   |   TiendaNova · Análisis de Rentabilidad Comercial   |   Pág. ")
fr.font.size = Pt(8)
fr.font.color.rgb = C_GRIS
fldChar1  = OxmlElement("w:fldChar"); fldChar1.set(qn("w:fldCharType"), "begin")
instrText = OxmlElement("w:instrText"); instrText.text = "PAGE"
fldChar2  = OxmlElement("w:fldChar"); fldChar2.set(qn("w:fldCharType"), "end")
fr._r.append(fldChar1); fr._r.append(instrText); fr._r.append(fldChar2)

# ── PORTADA ───────────────────────────────────────────────────────────────────
for _ in range(5):
    doc.add_paragraph()

p_tit = doc.add_paragraph()
p_tit.alignment = WD_ALIGN_PARAGRAPH.CENTER
rt = p_tit.add_run("Análisis de Rentabilidad Comercial")
rt.bold = True; rt.font.size = Pt(26); rt.font.color.rgb = C_AZUL_OSC

p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
rs = p_sub.add_run("Retail Multi-sucursal")
rs.font.size = Pt(18); rs.font.color.rgb = C_AZUL_MED

doc.add_paragraph()
linea_horizontal(doc)
doc.add_paragraph()

for linea, val in [
    ("Empresa analizada", "TiendaNova S.A. — Cadena minorista, 5 sucursales"),
    ("Período analizado", "Enero 2024 – Diciembre 2025 (24 meses)"),
    ("Herramienta",       "SQL (SQLite) — Esquema estrella, 5 consultas de negocio"),
    ("Analista",          "Luciano Villagrán — Analista de Datos"),
    ("Fecha",             date.today().strftime("%d de %B de %Y").replace(
        "January","enero").replace("February","febrero").replace("March","marzo")
        .replace("April","abril").replace("May","mayo").replace("June","junio")
        .replace("July","julio").replace("August","agosto").replace("September","septiembre")
        .replace("October","octubre").replace("November","noviembre").replace("December","diciembre")),
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p.add_run(f"{linea}: ")
    r1.bold = True; r1.font.size = Pt(11); r1.font.color.rgb = RGBColor(0x3D,0x3D,0x3D)
    r2 = p.add_run(val)
    r2.font.size = Pt(11); r2.font.color.rgb = C_GRIS

doc.add_page_break()

# ── SECCIÓN 1 — RESUMEN EJECUTIVO ─────────────────────────────────────────────
titulo_seccion(doc, "1. Resumen Ejecutivo")
linea_horizontal(doc)

p_box = doc.add_paragraph()
p_box.paragraph_format.left_indent  = Cm(0.5)
p_box.paragraph_format.space_before = Pt(4)
p_box.paragraph_format.space_after  = Pt(4)
set_para_bg(p_box, HEX_DIAG)
rb = p_box.add_run(
    "  Impacto total identificado: ~$2,0 M/año  +  ~$79.000 de capital liberado\n"
    "  Sobre un margen anual de ~$24,9 M → representa el 8 % de mejora de margen"
)
rb.bold = True; rb.font.size = Pt(12); rb.font.color.rgb = C_AZUL_OSC

doc.add_paragraph()

parrafo(doc,
    "El problema: TiendaNova facturaba $176,6 M en 24 meses con un margen promedio "
    "del 28,2 %, aparentemente saludable. Sin embargo, la rentabilidad se gestionaba 'a ojo': "
    "no existía visibilidad sobre qué productos se vendían a pérdida, qué clientes habían "
    "dejado de comprar ni por qué una sucursal rendía menos que el resto.",
    size=11, color=RGBColor(0x3D,0x3D,0x3D))

parrafo(doc,
    "La solución: modelé los datos en un esquema estrella (tabla de hechos ventas + "
    "3 dimensiones: clientes, productos y sucursales) y escribí 5 consultas SQL que detectan "
    "exactamente dónde se fuga el margen. Cada consulta responde una pregunta de negocio "
    "y termina con un número accionable.",
    size=11, color=RGBColor(0x3D,0x3D,0x3D), espacio_antes=4)

parrafo(doc,
    "El resultado: identifiqué $2,0 M/año de margen recuperable en 4 hallazgos accionables, "
    "sin incorporar un solo cliente nuevo y sin invertir en marketing. Solo dejando de "
    "perder lo que ya se estaba perdiendo.",
    size=11, color=RGBColor(0x3D,0x3D,0x3D), espacio_antes=4)

doc.add_page_break()

# ── SECCIÓN 2 — METODOLOGÍA Y DATOS ──────────────────────────────────────────
titulo_seccion(doc, "2. Metodología y Datos")
linea_horizontal(doc)

parrafo(doc, "Fuente de datos", size=11, bold=True, color=C_AZUL_OSC, espacio_antes=4)
parrafo(doc,
    "Base de datos SQLite generada sintéticamente con semilla fija (SEED=42), "
    "lo que garantiza reproducibilidad total: cualquier persona que ejecute "
    "generar_datos.py obtiene exactamente los mismos registros.",
    size=11, color=RGBColor(0x3D,0x3D,0x3D))

# Tabla de datasets
t1 = doc.add_table(rows=5, cols=3)
t1.style = "Table Grid"
encabezados_t1 = ["Tabla", "Registros", "Descripción"]
for i, h in enumerate(encabezados_t1):
    c = t1.rows[0].cells[i]
    c.text = h
    c.paragraphs[0].runs[0].bold = True
    c.paragraphs[0].runs[0].font.color.rgb = C_BLANCO
    set_cell_bg(c, HEX_AZUL_OSC)

filas_t1 = [
    ("ventas (hechos)", "15.700 líneas", "fecha · cantidad · precio · descuento · costo · ingreso · margen"),
    ("clientes",        "800 clientes",  "segmento · ciudad · región · fecha_alta"),
    ("productos",       "37 productos",  "categoría · costo_unitario · precio_lista"),
    ("sucursales",      "5 sucursales",  "nombre · ciudad · región"),
]
for i, (t, r, d) in enumerate(filas_t1):
    row = t1.rows[i+1].cells
    row[0].text = t; row[1].text = r; row[2].text = d
    bg = "F2F2F2" if i % 2 == 0 else "FFFFFF"
    for c in row:
        set_cell_bg(c, bg)

doc.add_paragraph()
parrafo(doc, "Herramientas y técnicas", size=11, bold=True, color=C_AZUL_OSC, espacio_antes=4)
for linea in [
    "• SQL (SQLite): 5 consultas con CTEs encadenadas (WITH), JOINs y métricas derivadas",
    "• Esquema estrella: tabla de hechos central + 3 dimensiones relacionadas",
    "• Validación previa: consulta de perfilado (00_perfilado.sql) para verificar integridad",
    "• Período: 24 meses (enero 2024 – diciembre 2025)",
]:
    parrafo(doc, linea, size=10.5, color=RGBColor(0x3D,0x3D,0x3D))

doc.add_page_break()

# ── SECCIÓN 3 — HALLAZGOS ─────────────────────────────────────────────────────
titulo_seccion(doc, "3. Hallazgos y Evidencia")
linea_horizontal(doc)

# ── H1: Margen negativo ───────────────────────────────────────────────────────
bloque_hallazgo(doc,
    numero=1,
    titulo="Productos vendidos a pérdida por descuentos agresivos",
    que_encontre=(
        "4 productos acumulan margen negativo durante los 24 meses analizados. "
        "El descuento promedio aplicado sobre estos productos fue del 17,6 %, "
        "suficiente para invertir el margen y convertir cada venta en una pérdida."
    ),
    accion="Quitar el descuento o repreciar estos 4 productos hasta margen ≥ 0 %.",
    impacto="~$0,91 M/año recuperables al eliminar el margen negativo.",
    supuesto="Margen negativo acumulado $1,82 M / 24 meses. Se supone mantenimiento "
             "de volumen; aun si baja algo al subir el precio, el resultado sigue siendo mejor "
             "que vender a pérdida."
)

# Tabla evidencia H1
t2 = doc.add_table(rows=5, cols=4)
t2.style = "Table Grid"
t2.paragraph_format = None
enc_t2 = ["Producto", "Categoría", "Facturación", "Margen total"]
for i, h in enumerate(enc_t2):
    c = t2.rows[0].cells[i]
    c.text = h
    c.paragraphs[0].runs[0].bold = True
    c.paragraphs[0].runs[0].font.color.rgb = C_BLANCO
    set_cell_bg(c, HEX_AZUL_MED)

datos_t2 = [
    ("P020", "Perfumería",  "$5.815.932", "-$730.785"),
    ("P017", "Limpieza",    "$4.046.601", "-$491.179"),
    ("P005", "Bebidas",     "$2.728.636", "-$365.970"),
    ("P009", "Almacén",     "$1.829.663", "-$235.225"),
]
for i, (p, cat, fac, mar) in enumerate(datos_t2):
    row = t2.rows[i+1].cells
    row[0].text = p; row[1].text = cat; row[2].text = fac
    row[3].text = mar
    set_cell_bg(row[3], HEX_ROJO_CL)
    row[3].paragraphs[0].runs[0].bold = True
    for j in range(3):
        set_cell_bg(row[j], "F2F2F2" if i % 2 == 0 else "FFFFFF")

p_pie = doc.add_paragraph()
p_pie.paragraph_format.left_indent = Cm(0.5)
rp = p_pie.add_run("Fuente: consulta 01_margen_negativo.sql sobre ventas 2024–2025.")
rp.font.size = Pt(9); rp.font.color.rgb = C_GRIS; rp.font.italic = True

# ── H2: Baja rotación ────────────────────────────────────────────────────────
bloque_hallazgo(doc,
    numero=2,
    titulo="Productos de baja rotación inmovilizan capital",
    que_encontre=(
        "5 productos (P036, P029, P023, P021, P003) registraron ≤ 1–5 unidades/mes "
        "en los últimos 6 meses del período analizado. Ocupan espacio en góndola y "
        "amarran capital que podría rotar en productos de mayor demanda."
    ),
    accion="Liquidar o descontinuar los 5 SKUs de baja rotación. Redirigir el espacio "
           "y el capital a los productos de mayor rotación.",
    impacto="~$79.000 de capital liberado (puntual) + ~$6.300/año de costo de oportunidad recuperado.",
    supuesto="Capital estimado asumiendo 3 meses de cobertura al ritmo actual de ventas. "
             "Costo de oportunidad calculado a tasa del 8 % anual sobre el capital inmovilizado."
)

# ── H3: Fuga de clientes ─────────────────────────────────────────────────────
bloque_hallazgo(doc,
    numero=3,
    titulo="156 clientes (21,7 %) no regresaron en 2025",
    que_encontre=(
        "De los 718 clientes activos en 2024, 156 (el 21,7 % de la base) no realizaron "
        "ninguna compra en 2025. Este grupo generó $2,45 M de margen en 2024, "
        "un valor que hoy está completamente perdido para la empresa."
    ),
    accion="Lanzar una campaña de recuperación priorizada (cupón dirigido, contacto comercial) "
           "sobre el listado de 156 clientes identificados, ordenados por margen aportado en 2024.",
    impacto="~$0,74 M/año al recuperar el 30 % del grupo perdido.",
    supuesto="Recuperación conservadora del 30 % de los 156 clientes. "
             "Ticket y frecuencia de compra similar a 2024."
)

doc.add_page_break()

# ── H4: Brecha sucursales ────────────────────────────────────────────────────
bloque_hallazgo(doc,
    numero=4,
    titulo="Sucursal Sur tiene el ticket 16 % por debajo de la red",
    que_encontre=(
        "La sucursal Sur (Bariloche) tiene un ticket promedio de $9.499, "
        "vs. una mediana de red de $11.267. La brecha es del 16 % y se mantiene "
        "constante en los 24 meses analizados, con un volumen de tickets sano (3.128), "
        "lo que descarta que sea un problema de tráfico."
    ),
    accion="Replicar en la sucursal Sur las prácticas comerciales de las sucursales líderes: "
           "estrategias de cross-sell, surtido ampliado en categorías de mayor valor "
           "y acciones para aumentar el tamaño promedio de compra.",
    impacto="~$0,39 M/año al cerrar el 50 % de la brecha de ticket (cierre total: ~$0,78 M/año).",
    supuesto="Cierre del 50 % de la brecha ($887 × 3.128 tickets × 28,2 % de margen). "
             "El volumen de tickets se mantiene constante."
)

# Gráfico sucursales
buf_suc = generar_grafico_sucursales()
doc.add_picture(buf_suc, width=Inches(5.8))
p_pie2 = doc.add_paragraph()
p_pie2.alignment = WD_ALIGN_PARAGRAPH.CENTER
rp2 = p_pie2.add_run(
    "Figura 1 — Ticket promedio por sucursal. La sucursal Sur registra el menor ticket "
    "de la red pese a un volumen de operaciones saludable. Fuente: 04_brecha_sucursales.sql."
)
rp2.font.size = Pt(9); rp2.font.color.rgb = C_GRIS; rp2.font.italic = True

doc.add_page_break()

# ── SECCIÓN 4 — CONCLUSIONES ──────────────────────────────────────────────────
titulo_seccion(doc, "4. Conclusiones y Próximos Pasos")
linea_horizontal(doc)

parrafo(doc,
    "El análisis reveló que el 8 % del margen anual de TiendaNova está siendo drenado "
    "por cuatro fuentes de pérdida identificables y corregibles con acciones comerciales "
    "de bajo a medio esfuerzo. Ninguna requiere nuevos clientes ni inversión en publicidad.",
    size=11, color=RGBColor(0x3D,0x3D,0x3D), espacio_antes=4)

# Tabla resumen de impacto
parrafo(doc, "Tabla resumen de impacto", size=11, bold=True, color=C_AZUL_OSC, espacio_antes=8)
t3 = doc.add_table(rows=6, cols=4)
t3.style = "Table Grid"
enc_t3 = ["Hallazgo", "Acción", "Impacto $/año", "Esfuerzo"]
for i, h in enumerate(enc_t3):
    c = t3.rows[0].cells[i]
    c.text = h
    c.paragraphs[0].runs[0].bold = True
    c.paragraphs[0].runs[0].font.color.rgb = C_BLANCO
    set_cell_bg(c, HEX_AZUL_OSC)

datos_t3 = [
    ("Productos bajo margen",      "Repreciar 4 productos",           "$0,91 M",  "Bajo"),
    ("Fuga de clientes",           "Campaña de recuperación",         "$0,74 M",  "Medio"),
    ("Brecha sucursal Sur",        "Replicar mejores prácticas",      "$0,39 M",  "Medio"),
    ("Capital inmovilizado",       "Liquidar 5 SKUs de baja rotación","$0,006 M + $79 k puntual", "Bajo"),
    ("TOTAL",                      "—",                               "~$2,0 M/año + $79 k", "—"),
]
for i, (h, a, imp, esf) in enumerate(datos_t3):
    row = t3.rows[i+1].cells
    row[0].text = h; row[1].text = a; row[2].text = imp; row[3].text = esf
    if i == 4:  # fila total
        for c in row:
            set_cell_bg(c, HEX_AZUL_CLA)
            c.paragraphs[0].runs[0].bold = True
    else:
        bg = "F2F2F2" if i % 2 == 0 else "FFFFFF"
        for c in row:
            set_cell_bg(c, bg)

# Caja impacto total
doc.add_paragraph()
p_tot = doc.add_paragraph()
p_tot.paragraph_format.space_before = Pt(6)
p_tot.paragraph_format.space_after  = Pt(6)
p_tot.paragraph_format.left_indent  = Cm(0.5)
set_para_bg(p_tot, HEX_DIAG)
rt2 = p_tot.add_run(
    "  Impacto total estimado del análisis: ~$2,0 M/año  +  ~$79.000 de capital liberado\n"
    "  Sobre un margen anual de ~$24,9 M → 8 % de mejora de margen identificada con 5 consultas SQL."
)
rt2.bold = True; rt2.font.size = Pt(11); rt2.font.color.rgb = C_AZUL_OSC

# Gráfico impacto
doc.add_paragraph()
buf_imp = generar_grafico_impacto()
doc.add_picture(buf_imp, width=Inches(5.8))
p_pie3 = doc.add_paragraph()
p_pie3.alignment = WD_ALIGN_PARAGRAPH.CENTER
rp3 = p_pie3.add_run(
    "Figura 2 — Impacto económico estimado por hallazgo. "
    "Los tres primeros hallazgos representan el 99,7 % del valor identificado."
)
rp3.font.size = Pt(9); rp3.font.color.rgb = C_GRIS; rp3.font.italic = True

doc.add_paragraph()
parrafo(doc, "Próximos pasos recomendados", size=11, bold=True, color=C_AZUL_OSC, espacio_antes=6)
for linea in [
    "1. Repreciar los 4 productos con margen negativo (bajo esfuerzo, mayor ROI inmediato — $0,91 M/año).",
    "2. Exportar el listado de 156 clientes fugados al equipo comercial para campaña de recuperación.",
    "3. Convocar reunión con el gerente de la sucursal Sur para revisar el modelo de cross-sell y surtido.",
]:
    parrafo(doc, linea, size=10.5, color=RGBColor(0x3D,0x3D,0x3D))

doc.add_paragraph()
linea_horizontal(doc)
p_cierre = doc.add_paragraph()
p_cierre.alignment = WD_ALIGN_PARAGRAPH.CENTER
rc = p_cierre.add_run(
    "Datos sintéticos con distribuciones realistas del sector retail  ·  "
    "Portfolio de Datos y Analítica  ·  Luciano Villagrán"
)
rc.font.size = Pt(9); rc.font.color.rgb = C_GRIS; rc.font.italic = True

# ══════════════════════════════════════════════════════════════════════════════
# GUARDAR DOCX Y EXPORTAR PDF
# ══════════════════════════════════════════════════════════════════════════════
doc.save(RUTA_DOCX)
print(f"DOCX generado: {RUTA_DOCX}")

try:
    from docx2pdf import convert
    convert(str(RUTA_DOCX), str(RUTA_PDF))
    print(f"PDF generado:  {RUTA_PDF}")
except Exception as e:
    print(f"PDF no generado: {e}")
    print("Abrí el DOCX en Word y guardá como PDF manualmente.")
